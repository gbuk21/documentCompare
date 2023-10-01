import docx
from nltk import tokenize
import db_ops as dops
conn=dops.connect()

old = docx.Document("/home/vasishk/Documents/old.docx")
new = docx.Document("/home/vasishk/Documents/new.docx")
all_old_paras = old.paragraphs
all_new_paras = new.paragraphs
len(all_new_paras)
len(all_old_paras)
new_sent_list=[]
old_sent_list=[]
new_matched_sent_list=[]
old_matched_sent_list=[]
for new_para in all_new_paras:
    #print(para.text)
    all_new_sentences=tokenize.sent_tokenize(new_para.text)
    #print(all_sentences)
    for new_sent in all_new_sentences:
        new_sent_list.append(new_sent) 
for old_para in all_old_paras:
    #print(para.text)
    all_old_sentences=tokenize.sent_tokenize(old_para.text)
    #print(all_sentences)
    for old_sent in all_old_sentences:
        old_sent_list.append(old_sent) 

old_sent_match_ind=0
new_sent_match_ind=0
loop_old_sent_ind=0
inserted_line_number=0
for i in new_sent_list:
    print('i  :'+i)
    print(loop_old_sent_ind)
    print(old_sent_match_ind)
    loop_old_sent_ind=0
    for j in old_sent_list:
        if loop_old_sent_ind<old_sent_match_ind:
           loop_old_sent_ind=loop_old_sent_ind+1
           continue
        else:
           loop_old_sent_ind=loop_old_sent_ind+1
        print('j  :'+j)
        if i==j or  old_sent_list[old_sent_match_ind+1]==new_sent_list[new_sent_match_ind+1]:
           inserted_line_number=inserted_line_number+1
           new_matched_sent_list.append(i)
           old_matched_sent_list.append(j)  
           dops.add_compare(conn,"","","","","",inserted_line_number,i,j)
           old_sent_match_ind=old_sent_match_ind+1
           new_sent_match_ind=new_sent_match_ind+1
           break
        else:
           old_sent_match_ind=old_sent_match_ind+1
           inserted_line_number=inserted_line_number+1
           new_matched_sent_list.append(i)
           old_matched_sent_list.append("")
           dops.add_compare(conn,"","","","","",inserted_line_number,i,"")
           break
print(old_matched_sent_list)
print(new_matched_sent_list)
